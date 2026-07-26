from __future__ import annotations

import csv
import difflib
import hashlib
import hmac
import io
import json
import os
import re
import secrets
import sqlite3
from collections import Counter, defaultdict
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any
from uuid import uuid4

import pdfplumber
import uvicorn
from docx import Document
from fastapi import Cookie, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, RedirectResponse, Response
from fastapi.staticfiles import StaticFiles
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

BASE = Path(__file__).resolve().parent
DATA = BASE / "data"
UPLOADS = DATA / "uploads"
DB_PATH = DATA / "vaelith.db"
DATA.mkdir(exist_ok=True)
UPLOADS.mkdir(exist_ok=True)

app = FastAPI(title="VAELITH LABS — Soluções em Engenharia", version="4.0-test")

ACCEPTED = {
    ".ifc", ".rvt", ".dwg", ".pdf", ".doc", ".docx", ".xls", ".xlsx", ".xlsm",
    ".csv", ".mpp", ".jpg", ".jpeg", ".png", ".webp", ".tif", ".tiff", ".json"
}

DISCIPLINE_KEYS = [
    ("ARQ", "Arquitetura"), ("ARQUIT", "Arquitetura"), ("ESTR", "Estrutura"),
    ("ELE", "Elétrica"), ("HID", "Hidráulica"), ("SANIT", "Sanitária"),
    ("INCEND", "Incêndio"), ("HVAC", "Climatização"), ("CLIMA", "Climatização"),
    ("ORC", "Orçamento"), ("ORÇ", "Orçamento"), ("CRONO", "Planejamento"),
    ("PLAN", "Planejamento"), ("INTERIOR", "Interiores"), ("FACH", "Fachada"),
    ("TELECOM", "Telecomunicações"), ("GAS", "Gás"), ("GÁS", "Gás")
]

ROOT_TYPES = {
    "IFCPROJECT", "IFCSITE", "IFCBUILDING", "IFCBUILDINGSTOREY", "IFCSPACE",
    "IFCWALL", "IFCWALLSTANDARDCASE", "IFCDOOR", "IFCWINDOW", "IFCSLAB", "IFCBEAM",
    "IFCCOLUMN", "IFCPIPESEGMENT", "IFCDUCTSEGMENT", "IFCCABLECARRIERSEGMENT",
    "IFCFLOWTERMINAL", "IFCFURNISHINGELEMENT", "IFCSTAIR", "IFCROOF", "IFCOPENINGELEMENT"
}


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def format_brl(value: float | int | None) -> str:
    if value is None:
        return "Não calculável"
    number = f"{float(value):,.2f}"
    number = number.replace(",", "_").replace(".", ",").replace("_", ".")
    return f"R$ {number}"


def conn() -> sqlite3.Connection:
    c = sqlite3.connect(DB_PATH)
    c.row_factory = sqlite3.Row
    c.execute("PRAGMA foreign_keys=ON")
    return c


def hash_password(password: str, salt: bytes | None = None) -> tuple[str, str]:
    salt = salt or secrets.token_bytes(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode(), salt, 180_000)
    return salt.hex(), digest.hex()


def verify_password(password: str, salt_hex: str, digest_hex: str) -> bool:
    _, digest = hash_password(password, bytes.fromhex(salt_hex))
    return hmac.compare_digest(digest, digest_hex)


def init_db() -> None:
    with conn() as c:
        c.executescript("""
        CREATE TABLE IF NOT EXISTS users(
            id TEXT PRIMARY KEY, name TEXT NOT NULL, email TEXT UNIQUE NOT NULL,
            salt TEXT NOT NULL, password_hash TEXT NOT NULL, created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS sessions(
            token TEXT PRIMARY KEY, user_id TEXT NOT NULL, expires_at TEXT NOT NULL,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
        );
        CREATE TABLE IF NOT EXISTS projects(
            id TEXT PRIMARY KEY, user_id TEXT NOT NULL, name TEXT NOT NULL,
            description TEXT NOT NULL DEFAULT '', created_at TEXT NOT NULL, updated_at TEXT NOT NULL,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
        );
        CREATE TABLE IF NOT EXISTS files(
            id TEXT PRIMARY KEY, project_id TEXT NOT NULL, original_name TEXT NOT NULL,
            stored_name TEXT NOT NULL, ext TEXT NOT NULL, size_bytes INTEGER NOT NULL,
            sha256 TEXT NOT NULL, discipline TEXT NOT NULL, revision TEXT NOT NULL,
            status TEXT NOT NULL, summary TEXT NOT NULL, details_json TEXT NOT NULL,
            uploaded_at TEXT NOT NULL,
            FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
        );
        CREATE TABLE IF NOT EXISTS changes(
            id TEXT PRIMARY KEY, project_id TEXT NOT NULL, code TEXT NOT NULL,
            title TEXT NOT NULL, request_text TEXT NOT NULL, location TEXT NOT NULL,
            element TEXT NOT NULL, stage TEXT NOT NULL, base_deadline TEXT,
            created_at TEXT NOT NULL, updated_at TEXT NOT NULL,
            FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
        );
        CREATE TABLE IF NOT EXISTS analyses(
            id TEXT PRIMARY KEY, project_id TEXT NOT NULL, change_id TEXT,
            result_json TEXT NOT NULL, created_at TEXT NOT NULL,
            FOREIGN KEY(project_id) REFERENCES projects(id) ON DELETE CASCADE
        );
        """)
        row = c.execute("SELECT id FROM users WHERE email=?", ("demo@vaelithlabs.com.br",)).fetchone()
        if not row:
            uid = uuid4().hex
            salt, ph = hash_password("vaelith")
            c.execute("INSERT INTO users VALUES(?,?,?,?,?,?)", (uid, "Usuário Demo", "demo@vaelithlabs.com.br", salt, ph, now_iso()))
            pid = uuid4().hex
            c.execute("INSERT INTO projects VALUES(?,?,?,?,?,?)", (
                pid, uid, "Projeto de demonstração", "Projeto vazio para testar arquivos reais.", now_iso(), now_iso()
            ))


init_db()


def require_user(session: str | None) -> sqlite3.Row:
    if not session:
        raise HTTPException(401, "Faça login para continuar")
    with conn() as c:
        row = c.execute("""
            SELECT u.* FROM sessions s JOIN users u ON u.id=s.user_id
            WHERE s.token=? AND s.expires_at>?
        """, (session, now_iso())).fetchone()
    if not row:
        raise HTTPException(401, "Sessão expirada")
    return row


def project_for_user(project_id: str, user_id: str) -> sqlite3.Row:
    with conn() as c:
        p = c.execute("SELECT * FROM projects WHERE id=? AND user_id=?", (project_id, user_id)).fetchone()
    if not p:
        raise HTTPException(404, "Projeto não encontrado")
    return p


def infer_revision(name: str) -> str:
    upper = name.upper()
    m = re.search(r"(?:^|[^A-Z0-9])(?:REV(?:IS[AÃ]O)?[ _.-]*|R)(\d{1,3})(?:[^A-Z0-9]|$)", upper)
    return f"R{int(m.group(1)):02d}" if m else "Não identificada"


def infer_discipline(name: str) -> str:
    upper = name.upper()
    for key, value in DISCIPLINE_KEYS:
        if key in upper:
            return value
    return "Não identificada"


def format_bytes(n: int) -> str:
    x = float(n)
    units = ["B", "KB", "MB", "GB"]
    i = 0
    while x >= 1024 and i < len(units)-1:
        x /= 1024
        i += 1
    return f"{x:.1f} {units[i]}" if i else f"{int(x)} B"


def normalize_ifc_line(line: str) -> str:
    return re.sub(r"\s+", "", line).upper()


def parse_ifc(path: Path) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8", errors="replace")
    schema_m = re.search(r"FILE_SCHEMA\s*\(\s*\(\s*'([^']+)'", text, re.I)
    entity_counts = Counter(re.findall(r"=\s*(IFC[A-Z0-9_]+)\s*\(", text, re.I))
    elements: dict[str, dict[str, Any]] = {}
    duplicate_guids: list[str] = []
    for m in re.finditer(r"#(\d+)\s*=\s*(IFC[A-Z0-9_]+)\s*\((.*?)\);", text, re.I | re.S):
        express_id, entity, args = m.group(1), m.group(2).upper(), m.group(3)
        if entity not in ROOT_TYPES:
            continue
        str_args = re.findall(r"'((?:[^']|'')*)'", args)
        guid = str_args[0] if str_args else f"EXPRESS-{express_id}"
        name = str_args[2] if len(str_args) > 2 else (str_args[1] if len(str_args) > 1 else "")
        rec = {
            "guid": guid,
            "expressId": int(express_id),
            "entity": entity,
            "name": name,
            "signature": hashlib.sha1(normalize_ifc_line(m.group(0)).encode()).hexdigest(),
        }
        if guid in elements:
            duplicate_guids.append(guid)
        elements[guid] = rec
    return {
        "schema": schema_m.group(1) if schema_m else None,
        "entityCounts": entity_counts.most_common(40),
        "entityTotal": sum(entity_counts.values()),
        "elements": elements,
        "elementCount": len(elements),
        "duplicateGuids": duplicate_guids[:100],
    }


def find_header_map(rows: list[list[Any]]) -> tuple[int | None, dict[str, int]]:
    aliases = {
        "description": ["descricao", "descrição", "servico", "serviço", "item", "atividade"],
        "quantity": ["quantidade", "qtd", "quant."],
        "unit": ["unidade", "un"],
        "unit_price": ["preco unitario", "preço unitário", "valor unitario", "valor unitário", "pu"],
        "total": ["total", "valor total", "subtotal"],
        "start": ["inicio", "início", "data inicio", "data início"],
        "end": ["fim", "termino", "término", "data fim"],
        "duration": ["duracao", "duração", "dias", "prazo"],
        "predecessor": ["predecessora", "predecessor", "dependencia", "dependência"],
    }
    for idx, row in enumerate(rows[:25]):
        normalized = [re.sub(r"\s+", " ", str(v or "").strip().lower()) for v in row]
        mapping: dict[str, int] = {}
        for key, names in aliases.items():
            for col, val in enumerate(normalized):
                if any(alias == val or alias in val for alias in names):
                    mapping[key] = col
                    break
        if "description" in mapping and ("quantity" in mapping or "duration" in mapping or "start" in mapping):
            return idx, mapping
    return None, {}


def parse_workbook(path: Path) -> dict[str, Any]:
    wb = load_workbook(path, read_only=True, data_only=True)
    sheets: list[dict[str, Any]] = []
    budgets: list[dict[str, Any]] = []
    schedules: list[dict[str, Any]] = []
    for ws in wb.worksheets[:30]:
        rows = [[c for c in row] for row in ws.iter_rows(values_only=True)]
        header_idx, mapping = find_header_map(rows)
        preview = [[str(v or "")[:100] for v in row[:10]] for row in rows[:8]]
        sheets.append({"name": ws.title, "rows": ws.max_row, "columns": ws.max_column, "preview": preview})
        if header_idx is not None:
            for row in rows[header_idx+1:]:
                desc = str(row[mapping["description"]] or "").strip() if mapping.get("description") is not None and mapping["description"] < len(row) else ""
                if not desc:
                    continue
                if "quantity" in mapping or "unit_price" in mapping or "total" in mapping:
                    qty = row[mapping["quantity"]] if "quantity" in mapping and mapping["quantity"] < len(row) else None
                    pu = row[mapping["unit_price"]] if "unit_price" in mapping and mapping["unit_price"] < len(row) else None
                    total = row[mapping["total"]] if "total" in mapping and mapping["total"] < len(row) else None
                    budgets.append({
                        "sheet": ws.title, "description": desc,
                        "quantity": qty, "unit": row[mapping["unit"]] if "unit" in mapping and mapping["unit"] < len(row) else "",
                        "unitPrice": pu, "total": total
                    })
                if "duration" in mapping or "start" in mapping:
                    schedules.append({
                        "sheet": ws.title, "activity": desc,
                        "duration": row[mapping["duration"]] if "duration" in mapping and mapping["duration"] < len(row) else None,
                        "start": str(row[mapping["start"]]) if "start" in mapping and mapping["start"] < len(row) and row[mapping["start"]] is not None else None,
                        "end": str(row[mapping["end"]]) if "end" in mapping and mapping["end"] < len(row) and row[mapping["end"]] is not None else None,
                        "predecessor": row[mapping["predecessor"]] if "predecessor" in mapping and mapping["predecessor"] < len(row) else None,
                    })
    return {"sheets": sheets, "budgetRows": budgets[:5000], "scheduleRows": schedules[:5000]}


def parse_csv(path: Path) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8-sig", errors="replace")
    sample = text[:8000]
    dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|") if sample.strip() else csv.excel
    rows = list(csv.reader(io.StringIO(text), dialect))
    header_idx, mapping = find_header_map(rows)
    result = {"preview": rows[:20], "budgetRows": [], "scheduleRows": []}
    if header_idx is not None:
        for row in rows[header_idx+1:]:
            if not row:
                continue
            desc = row[mapping["description"]].strip() if mapping["description"] < len(row) else ""
            if not desc:
                continue
            if "quantity" in mapping or "unit_price" in mapping or "total" in mapping:
                result["budgetRows"].append({
                    "description": desc,
                    "quantity": row[mapping["quantity"]] if "quantity" in mapping and mapping["quantity"] < len(row) else None,
                    "unit": row[mapping["unit"]] if "unit" in mapping and mapping["unit"] < len(row) else "",
                    "unitPrice": row[mapping["unit_price"]] if "unit_price" in mapping and mapping["unit_price"] < len(row) else None,
                    "total": row[mapping["total"]] if "total" in mapping and mapping["total"] < len(row) else None,
                })
            if "duration" in mapping or "start" in mapping:
                result["scheduleRows"].append({
                    "activity": desc,
                    "duration": row[mapping["duration"]] if "duration" in mapping and mapping["duration"] < len(row) else None,
                    "start": row[mapping["start"]] if "start" in mapping and mapping["start"] < len(row) else None,
                    "end": row[mapping["end"]] if "end" in mapping and mapping["end"] < len(row) else None,
                    "predecessor": row[mapping["predecessor"]] if "predecessor" in mapping and mapping["predecessor"] < len(row) else None,
                })
    return result


def analyze_file(path: Path, original: str) -> tuple[str, str, dict[str, Any]]:
    ext = path.suffix.lower()
    try:
        if ext == ".ifc":
            d = parse_ifc(path)
            return "Processado", f"IFC {d.get('schema') or ''}: {d['elementCount']} elementos e {d['entityTotal']} entidades.", d
        if ext in {".xlsx", ".xlsm"}:
            d = parse_workbook(path)
            return "Processado", f"{len(d['sheets'])} abas; {len(d['budgetRows'])} linhas de orçamento e {len(d['scheduleRows'])} atividades detectadas.", d
        if ext == ".csv":
            d = parse_csv(path)
            return "Processado", f"CSV: {len(d['budgetRows'])} linhas de orçamento e {len(d['scheduleRows'])} atividades detectadas.", d
        if ext == ".pdf":
            parts = []
            with pdfplumber.open(path) as pdf:
                count = len(pdf.pages)
                for page in pdf.pages[:25]:
                    parts.append(page.extract_text() or "")
            text = "\n".join(parts)
            return "Processado", f"PDF com {count} páginas; {len(text)} caracteres extraídos.", {"pages": count, "text": text[:100000]}
        if ext == ".docx":
            doc = Document(path)
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            tables = []
            for table in doc.tables[:20]:
                tables.append([[cell.text for cell in row.cells] for row in table.rows[:100]])
            return "Processado", f"Word com {len(paras)} parágrafos e {len(doc.tables)} tabelas.", {"paragraphs": paras[:3000], "tables": tables}
        if ext in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}:
            with Image.open(path) as im:
                return "Processado", f"Imagem {im.width} × {im.height}px.", {"width": im.width, "height": im.height, "mode": im.mode}
        if ext in {".rvt", ".dwg"}:
            return "Conversão necessária", "Arquivo aceito. Para maquete e clash detection, exporte também para IFC.", {"recommended": "IFC", "nativeGeometry": False}
        if ext == ".mpp":
            return "Conversão necessária", "Arquivo aceito. Exporte o cronograma para XLSX ou CSV para cálculo nesta versão.", {"recommended": "XLSX/CSV"}
        if ext in {".doc", ".xls"}:
            return "Conversão necessária", f"Arquivo legado aceito. Converta para {'.docx' if ext == '.doc' else '.xlsx'} para leitura automática.", {}
        return "Catalogado", "Arquivo salvo para referência.", {}
    except Exception as exc:
        return "Erro de processamento", f"O arquivo foi salvo, mas a leitura falhou: {type(exc).__name__}.", {"error": str(exc)[:1000]}


def file_rows(project_id: str) -> list[dict[str, Any]]:
    with conn() as c:
        rows = c.execute("SELECT * FROM files WHERE project_id=? ORDER BY uploaded_at", (project_id,)).fetchall()
    out = []
    for r in rows:
        d = dict(r)
        d["details"] = json.loads(d.pop("details_json"))
        d["size"] = format_bytes(d["size_bytes"])
        d["downloadUrl"] = f"/api/files/{d['id']}/raw"
        out.append(d)
    return out


def revision_number(rev: str) -> int:
    m = re.search(r"(\d+)", rev or "")
    return int(m.group(1)) if m else -1


def compare_ifc_files(a: dict[str, Any], b: dict[str, Any]) -> dict[str, Any]:
    ea = a.get("details", {}).get("elements", {})
    eb = b.get("details", {}).get("elements", {})
    ga, gb = set(ea), set(eb)
    added = [eb[g] for g in sorted(gb-ga)]
    removed = [ea[g] for g in sorted(ga-gb)]
    modified = []
    for g in sorted(ga & gb):
        if ea[g].get("signature") != eb[g].get("signature"):
            modified.append({"before": ea[g], "after": eb[g]})
    return {
        "discipline": b["discipline"], "from": a["revision"], "to": b["revision"],
        "fileFrom": a["original_name"], "fileTo": b["original_name"],
        "added": added[:1000], "removed": removed[:1000], "modified": modified[:1000],
        "counts": {"added": len(added), "removed": len(removed), "modified": len(modified)}
    }


def safe_float(v: Any) -> float | None:
    if v is None or v == "":
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip().replace("R$", "").replace(" ", "")
    if "," in s and "." in s:
        s = s.replace(".", "").replace(",", ".")
    elif "," in s:
        s = s.replace(",", ".")
    try:
        return float(s)
    except Exception:
        return None


def normalize_words(text: str) -> set[str]:
    stop = {"de", "da", "do", "e", "em", "para", "com", "um", "uma", "o", "a", "os", "as", "no", "na"}
    words = re.findall(r"[a-záàâãéêíóôõúç0-9]+", text.lower())
    return {w for w in words if len(w) > 2 and w not in stop}


def expand_change_words(words: set[str]) -> set[str]:
    expanded = set(words)
    groups = [
        ({"porta", "vão", "vao", "abertura"}, {"alvenaria", "vãos", "vaos", "demolição", "demolicao", "reboco", "pintura", "acabamento", "porta"}),
        ({"interruptor", "tomada", "elétrico", "eletrico", "elétrica", "eletrica", "eletroduto"}, {"instalações", "instalacoes", "elétrica", "eletrica", "remanejamento", "interruptor"}),
        ({"tubulação", "tubulacao", "hidráulica", "hidraulica", "esgoto"}, {"hidráulica", "hidraulica", "tubulação", "tubulacao", "instalações", "instalacoes"}),
        ({"parede", "alvenaria"}, {"parede", "alvenaria", "reboco", "pintura", "demolição", "demolicao"}),
        ({"forro", "climatização", "climatizacao", "duto"}, {"forro", "climatização", "climatizacao", "duto", "acabamento"}),
    ]
    for triggers, additions in groups:
        if words & triggers:
            expanded |= additions
    return expanded


def match_budget(change: sqlite3.Row | None, files: list[dict[str, Any]]) -> dict[str, Any]:
    budget_rows: list[dict[str, Any]] = []
    for f in files:
        budget_rows.extend(f.get("details", {}).get("budgetRows", []))
    if not budget_rows:
        return {"status": "not_calculable", "reason": "Nenhuma planilha de orçamento com colunas reconhecidas foi vinculada.", "matches": [], "total": None}
    query = " ".join([change["title"], change["request_text"], change["element"], change["location"]]) if change else ""
    qwords = expand_change_words(normalize_words(query))
    scored = []
    for row in budget_rows:
        desc = str(row.get("description", ""))
        words = normalize_words(desc)
        score = len(qwords & words)
        if score:
            qty = safe_float(row.get("quantity"))
            pu = safe_float(row.get("unitPrice"))
            total = safe_float(row.get("total"))
            if total is None and qty is not None and pu is not None:
                total = qty * pu
            scored.append({**row, "matchScore": score, "calculatedTotal": total})
    scored.sort(key=lambda r: (-r["matchScore"], str(r.get("description"))))
    matches = scored[:30]
    values = [r["calculatedTotal"] for r in matches if r.get("calculatedTotal") is not None]
    return {
        "status": "calculated" if values else "partial",
        "reason": "Itens semelhantes localizados no orçamento. A seleção deve ser validada pelo orçamentista." if matches else "Nenhum item do orçamento correspondeu ao texto da mudança.",
        "matches": matches,
        "total": sum(values) if values else None,
    }


def match_schedule(change: sqlite3.Row | None, files: list[dict[str, Any]]) -> dict[str, Any]:
    rows = []
    for f in files:
        rows.extend(f.get("details", {}).get("scheduleRows", []))
    if not rows:
        return {"status": "not_calculable", "reason": "Nenhum cronograma XLSX/CSV reconhecido foi vinculado.", "matches": [], "days": None}
    query = " ".join([change["title"], change["request_text"], change["element"], change["location"]]) if change else ""
    qwords = expand_change_words(normalize_words(query))
    scored = []
    for row in rows:
        activity = str(row.get("activity", ""))
        score = len(qwords & normalize_words(activity))
        if score:
            duration = safe_float(row.get("duration"))
            scored.append({**row, "matchScore": score, "durationNumber": duration})
    scored.sort(key=lambda r: (-r["matchScore"], str(r.get("activity"))))
    matches = scored[:30]
    durations = [r["durationNumber"] for r in matches if r.get("durationNumber") is not None]
    return {
        "status": "calculated" if durations else "partial",
        "reason": "Atividades relacionadas foram encontradas; o impacto líquido depende das predecessoras e da situação executada." if matches else "Nenhuma atividade do cronograma correspondeu à mudança.",
        "matches": matches,
        "days": max(durations) if durations else None,
    }


def run_analysis(project_id: str, change_id: str | None = None) -> dict[str, Any]:
    files = file_rows(project_id)
    with conn() as c:
        change = c.execute("SELECT * FROM changes WHERE id=? AND project_id=?", (change_id, project_id)).fetchone() if change_id else c.execute("SELECT * FROM changes WHERE project_id=? ORDER BY created_at DESC LIMIT 1", (project_id,)).fetchone()
    issues: list[dict[str, Any]] = []
    by_disc: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for f in files:
        by_disc[f["discipline"]].append(f)
        if f["status"] in {"Erro de processamento", "Conversão necessária"}:
            issues.append({"severity": "warning", "category": "Arquivo", "title": f["original_name"], "detail": f["summary"], "fileId": f["id"]})
        if f["revision"] == "Não identificada":
            issues.append({"severity": "warning", "category": "Revisão", "title": "Revisão não identificada", "detail": f"Informe a revisão do arquivo {f['original_name']}.", "fileId": f["id"]})
        if f["discipline"] == "Não identificada":
            issues.append({"severity": "warning", "category": "Disciplina", "title": "Disciplina não identificada", "detail": f"Classifique o arquivo {f['original_name']}.", "fileId": f["id"]})
    latest = {d: max((revision_number(f["revision"]) for f in fs), default=-1) for d, fs in by_disc.items() if d != "Não identificada"}
    if latest:
        top = max(latest.values())
        for d, rev in latest.items():
            if rev >= 0 and rev < top:
                issues.append({"severity": "high", "category": "Compatibilidade de versões", "title": f"{d} está desatualizado", "detail": f"Última revisão R{rev:02d}; revisão mais alta do projeto R{top:02d}."})
    ifc_comparisons = []
    for disc, fs in by_disc.items():
        ifcs = [f for f in fs if f["ext"].lower() == ".ifc" and f["status"] == "Processado"]
        ifcs.sort(key=lambda f: revision_number(f["revision"]))
        for a, b in zip(ifcs, ifcs[1:]):
            comp = compare_ifc_files(a, b)
            ifc_comparisons.append(comp)
            if comp["counts"]["added"] or comp["counts"]["removed"] or comp["counts"]["modified"]:
                issues.append({
                    "severity": "info", "category": "Comparação IFC",
                    "title": f"{disc}: {a['revision']} → {b['revision']}",
                    "detail": f"{comp['counts']['added']} adicionados, {comp['counts']['removed']} removidos e {comp['counts']['modified']} modificados."
                })
    duplicate_hashes = defaultdict(list)
    for f in files:
        duplicate_hashes[f["sha256"]].append(f)
    for same in duplicate_hashes.values():
        if len(same) > 1:
            issues.append({"severity": "info", "category": "Duplicidade", "title": "Arquivos idênticos", "detail": ", ".join(x["original_name"] for x in same)})
    text_comparisons = []
    for disc, fs in by_disc.items():
        docs = [f for f in fs if f["ext"].lower() in {".pdf", ".docx"} and f["status"] == "Processado"]
        docs.sort(key=lambda f: revision_number(f["revision"]))
        for a, b in zip(docs, docs[1:]):
            ta = a["details"].get("text", "\n".join(a["details"].get("paragraphs", [])))
            tb = b["details"].get("text", "\n".join(b["details"].get("paragraphs", [])))
            ratio = difflib.SequenceMatcher(None, ta[:100000], tb[:100000]).ratio() if ta or tb else 1.0
            text_comparisons.append({"discipline": disc, "from": a["revision"], "to": b["revision"], "similarity": round(ratio*100, 1), "fileFrom": a["original_name"], "fileTo": b["original_name"]})
    budget = match_budget(change, files)
    schedule = match_schedule(change, files)
    if budget["status"] == "not_calculable":
        issues.append({"severity": "warning", "category": "Custo", "title": "Impacto financeiro não calculável", "detail": budget["reason"]})
    if schedule["status"] == "not_calculable":
        issues.append({"severity": "warning", "category": "Prazo", "title": "Impacto no cronograma não calculável", "detail": schedule["reason"]})
    ifc_files = [f for f in files if f["ext"].lower() == ".ifc" and f["status"] == "Processado"]
    geometric = {
        "status": "ready" if len(ifc_files) >= 2 else "insufficient",
        "reason": "Carregue ao menos dois modelos IFC de disciplinas diferentes para executar o pré-clash 3D." if len(ifc_files) < 2 else "Os modelos estão disponíveis para visualização e pré-clash por envelopes no navegador.",
        "ifcFiles": [{"id": f["id"], "name": f["original_name"], "discipline": f["discipline"], "revision": f["revision"], "url": f"/api/files/{f['id']}/raw"} for f in ifc_files]
    }
    score = 0
    score += 20 if ifc_files else 0
    score += 20 if budget["status"] != "not_calculable" else 0
    score += 20 if schedule["status"] != "not_calculable" else 0
    score += 20 if all(f["revision"] != "Não identificada" for f in files) and files else 0
    score += 20 if all(f["discipline"] != "Não identificada" for f in files) and files else 0
    conclusion_parts = []
    conclusion_parts.append(f"Foram processados {len(files)} arquivos em {len(by_disc)} grupos de disciplina.")
    if ifc_comparisons:
        csum = {k: sum(c["counts"][k] for c in ifc_comparisons) for k in ["added", "removed", "modified"]}
        conclusion_parts.append(f"Nas revisões IFC, foram identificados {csum['added']} elementos adicionados, {csum['removed']} removidos e {csum['modified']} modificados.")
    if budget["total"] is not None:
        conclusion_parts.append(f"Os itens semelhantes localizados no orçamento somam {format_brl(budget['total'])}, valor que precisa de validação antes de ser tratado como impacto da mudança.")
    else:
        conclusion_parts.append("O custo ainda não pode ser fechado com fidelidade porque não há correspondência orçamentária suficiente.")
    if schedule["days"] is not None:
        conclusion_parts.append(f"A maior duração entre as atividades relacionadas é de {schedule['days']:g} dia(s), mas o impacto líquido depende das predecessoras e do estágio da obra.")
    else:
        conclusion_parts.append("O prazo ainda não pode ser recalculado com fidelidade porque não há cronograma estruturado ou atividades correspondentes.")
    result = {
        "id": uuid4().hex, "createdAt": now_iso(), "projectId": project_id,
        "change": dict(change) if change else None,
        "files": files, "issues": issues, "revisionMatrix": latest,
        "ifcComparisons": ifc_comparisons, "textComparisons": text_comparisons,
        "budget": budget, "schedule": schedule, "geometric": geometric,
        "dataQuality": score, "conclusion": " ".join(conclusion_parts),
        "trace": [
            {"step": 1, "title": "Inventário", "detail": f"{len(files)} arquivos recebidos e classificados por formato, disciplina e revisão."},
            {"step": 2, "title": "Revisões", "detail": f"{len(ifc_comparisons)} comparações IFC e {len(text_comparisons)} comparações documentais executadas."},
            {"step": 3, "title": "Compatibilidade", "detail": f"{len(issues)} alertas de versões, dados, conversão ou duplicidade foram gerados."},
            {"step": 4, "title": "Custo", "detail": budget["reason"]},
            {"step": 5, "title": "Prazo", "detail": schedule["reason"]},
            {"step": 6, "title": "Geometria", "detail": geometric["reason"]},
        ]
    }
    with conn() as c:
        c.execute("INSERT INTO analyses VALUES(?,?,?,?,?)", (result["id"], project_id, change["id"] if change else None, json.dumps(result, ensure_ascii=False), result["createdAt"]))
    return result


def latest_analysis(project_id: str) -> dict[str, Any] | None:
    with conn() as c:
        row = c.execute("SELECT result_json FROM analyses WHERE project_id=? ORDER BY created_at DESC LIMIT 1", (project_id,)).fetchone()
    return json.loads(row[0]) if row else None


@app.get("/api/health")
def health():
    return {"ok": True, "version": app.version, "time": now_iso()}


@app.post("/api/auth/login")
def login(payload: dict[str, Any]):
    email = str(payload.get("email", "")).strip().lower()
    password = str(payload.get("password", ""))
    with conn() as c:
        user = c.execute("SELECT * FROM users WHERE email=?", (email,)).fetchone()
        if not user or not verify_password(password, user["salt"], user["password_hash"]):
            raise HTTPException(401, "E-mail ou senha incorretos")
        token = secrets.token_urlsafe(32)
        expires = (datetime.now(timezone.utc) + timedelta(days=7)).isoformat()
        c.execute("INSERT INTO sessions VALUES(?,?,?)", (token, user["id"], expires))
    resp = JSONResponse({"ok": True, "user": {"id": user["id"], "name": user["name"], "email": user["email"]}})
    resp.set_cookie("vaelith_session", token, httponly=True, samesite="lax", max_age=7*86400)
    return resp


@app.post("/api/auth/logout")
def logout(vaelith_session: str | None = Cookie(default=None)):
    if vaelith_session:
        with conn() as c:
            c.execute("DELETE FROM sessions WHERE token=?", (vaelith_session,))
    resp = JSONResponse({"ok": True})
    resp.delete_cookie("vaelith_session")
    return resp


@app.get("/api/auth/me")
def me(vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    return {"id": u["id"], "name": u["name"], "email": u["email"]}


@app.get("/api/projects")
def list_projects(vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    with conn() as c:
        rows = c.execute("SELECT * FROM projects WHERE user_id=? ORDER BY updated_at DESC", (u["id"],)).fetchall()
    return [dict(r) for r in rows]


@app.post("/api/projects")
def create_project(payload: dict[str, Any], vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    name = str(payload.get("name", "")).strip()
    if not name:
        raise HTTPException(400, "Informe o nome do projeto")
    pid = uuid4().hex
    with conn() as c:
        c.execute("INSERT INTO projects VALUES(?,?,?,?,?,?)", (pid, u["id"], name, str(payload.get("description", "")), now_iso(), now_iso()))
    return {"id": pid, "name": name}


@app.get("/api/projects/{project_id}/state")
def project_state(project_id: str, vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    p = project_for_user(project_id, u["id"])
    with conn() as c:
        changes = [dict(r) for r in c.execute("SELECT * FROM changes WHERE project_id=? ORDER BY created_at DESC", (project_id,)).fetchall()]
    return {"project": dict(p), "files": file_rows(project_id), "changes": changes, "analysis": latest_analysis(project_id)}


@app.post("/api/projects/{project_id}/changes")
def create_change(project_id: str, payload: dict[str, Any], vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    project_for_user(project_id, u["id"])
    cid = uuid4().hex
    code = str(payload.get("code") or f"SM-{datetime.now().strftime('%Y%m%d-%H%M')}")
    values = (
        cid, project_id, code, str(payload.get("title", "Mudança sem título")),
        str(payload.get("requestText", "")), str(payload.get("location", "")),
        str(payload.get("element", "")), str(payload.get("stage", "Não informado")),
        str(payload.get("baseDeadline", "")) or None, now_iso(), now_iso()
    )
    with conn() as c:
        c.execute("INSERT INTO changes VALUES(?,?,?,?,?,?,?,?,?,?,?)", values)
    return {"id": cid, "code": code}


@app.post("/api/projects/{project_id}/files")
async def upload_files(
    project_id: str,
    files: list[UploadFile] = File(...),
    discipline: str = Form(default=""), revision: str = Form(default=""),
    vaelith_session: str | None = Cookie(default=None)
):
    u = require_user(vaelith_session)
    project_for_user(project_id, u["id"])
    added = []
    for upload in files:
        original = Path(upload.filename or "arquivo").name
        ext = Path(original).suffix.lower()
        if ext not in ACCEPTED:
            raise HTTPException(415, f"Formato não aceito: {ext or 'sem extensão'}")
        fid = uuid4().hex
        stored_name = f"{fid}{ext}"
        target = UPLOADS / stored_name
        sha = hashlib.sha256()
        size = 0
        with target.open("wb") as out:
            while chunk := await upload.read(1024 * 1024):
                sha.update(chunk); size += len(chunk); out.write(chunk)
        status, summary, details = analyze_file(target, original)
        disc = discipline.strip() or infer_discipline(original)
        rev = revision.strip().upper() or infer_revision(original)
        with conn() as c:
            c.execute("INSERT INTO files VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)", (
                fid, project_id, original, stored_name, ext, size, sha.hexdigest(), disc, rev,
                status, summary, json.dumps(details, ensure_ascii=False), now_iso()
            ))
            c.execute("UPDATE projects SET updated_at=? WHERE id=?", (now_iso(), project_id))
        added.append({"id": fid, "name": original, "status": status, "summary": summary, "discipline": disc, "revision": rev})
    return {"ok": True, "added": added, "files": file_rows(project_id)}


@app.patch("/api/files/{file_id}")
def update_file(file_id: str, payload: dict[str, Any], vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    with conn() as c:
        row = c.execute("SELECT f.*,p.user_id FROM files f JOIN projects p ON p.id=f.project_id WHERE f.id=?", (file_id,)).fetchone()
        if not row or row["user_id"] != u["id"]:
            raise HTTPException(404, "Arquivo não encontrado")
        c.execute("UPDATE files SET discipline=?, revision=? WHERE id=?", (
            str(payload.get("discipline", row["discipline"])), str(payload.get("revision", row["revision"])).upper(), file_id
        ))
    return {"ok": True}


@app.delete("/api/files/{file_id}")
def delete_file(file_id: str, vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    with conn() as c:
        row = c.execute("SELECT f.*,p.user_id FROM files f JOIN projects p ON p.id=f.project_id WHERE f.id=?", (file_id,)).fetchone()
        if not row or row["user_id"] != u["id"]:
            raise HTTPException(404, "Arquivo não encontrado")
        (UPLOADS / row["stored_name"]).unlink(missing_ok=True)
        c.execute("DELETE FROM files WHERE id=?", (file_id,))
    return {"ok": True}


@app.get("/api/files/{file_id}/raw")
def raw_file(file_id: str, vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    with conn() as c:
        row = c.execute("SELECT f.*,p.user_id FROM files f JOIN projects p ON p.id=f.project_id WHERE f.id=?", (file_id,)).fetchone()
    if not row or row["user_id"] != u["id"]:
        raise HTTPException(404, "Arquivo não encontrado")
    return FileResponse(UPLOADS / row["stored_name"], filename=row["original_name"])


@app.post("/api/projects/{project_id}/analyze")
def analyze_project(project_id: str, payload: dict[str, Any] | None = None, vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    project_for_user(project_id, u["id"])
    return run_analysis(project_id, (payload or {}).get("changeId"))


@app.get("/api/projects/{project_id}/export/{fmt}")
def export(project_id: str, fmt: str, vaelith_session: str | None = Cookie(default=None)):
    u = require_user(vaelith_session)
    p = project_for_user(project_id, u["id"])
    analysis = latest_analysis(project_id) or run_analysis(project_id)
    fmt = fmt.lower()
    title = f"VAELITH — {p['name']}"
    if fmt == "json":
        return Response(json.dumps(analysis, ensure_ascii=False, indent=2), media_type="application/json", headers={"Content-Disposition": "attachment; filename=vaelith-analise.json"})
    if fmt == "xlsx":
        wb = Workbook(); ws = wb.active; ws.title = "Resumo"
        rows = [["VAELITH LABS", "Soluções em Engenharia"], ["Projeto", p["name"]], ["Qualidade dos dados (%)", analysis["dataQuality"]], ["Conclusão", analysis["conclusion"]]]
        for r in rows: ws.append(r)
        ws.column_dimensions["A"].width = 30; ws.column_dimensions["B"].width = 110
        ws["A1"].fill = ws["B1"].fill = PatternFill("solid", fgColor="CFFF47"); ws["A1"].font = Font(bold=True)
        for name, headers, data in [
            ("Alertas", ["Severidade","Categoria","Título","Detalhe"], [[i.get("severity"),i.get("category"),i.get("title"),i.get("detail")] for i in analysis["issues"]]),
            ("Comparações IFC", ["Disciplina","De","Para","Adicionados","Removidos","Modificados"], [[x["discipline"],x["from"],x["to"],x["counts"]["added"],x["counts"]["removed"],x["counts"]["modified"]] for x in analysis["ifcComparisons"]]),
            ("Orçamento", ["Descrição","Quantidade","Unidade","Preço unitário","Total","Pontuação"], [[x.get("description"),x.get("quantity"),x.get("unit"),x.get("unitPrice"),x.get("calculatedTotal"),x.get("matchScore")] for x in analysis["budget"]["matches"]]),
            ("Cronograma", ["Atividade","Duração","Início","Fim","Predecessora","Pontuação"], [[x.get("activity"),x.get("duration"),x.get("start"),x.get("end"),x.get("predecessor"),x.get("matchScore")] for x in analysis["schedule"]["matches"]]),
        ]:
            sh = wb.create_sheet(name); sh.append(headers)
            for row in data: sh.append(row)
            for cell in sh[1]: cell.fill = PatternFill("solid", fgColor="CFFF47"); cell.font = Font(bold=True)
            for col in sh.columns: sh.column_dimensions[col[0].column_letter].width = min(60, max(12, max(len(str(c.value or "")) for c in col)+2))
            for row in sh.iter_rows():
                for cell in row: cell.alignment = Alignment(vertical="top", wrap_text=True)
        out = io.BytesIO(); wb.save(out)
        return Response(out.getvalue(), media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition":"attachment; filename=vaelith-analise.xlsx"})
    if fmt == "docx":
        doc = Document(); doc.add_heading("VAELITH LABS", 0); doc.add_paragraph("Soluções em Engenharia")
        doc.add_heading(p["name"], 1); doc.add_heading("Conclusão", 2); doc.add_paragraph(analysis["conclusion"])
        doc.add_heading("Linha de raciocínio", 2)
        for s in analysis["trace"]: doc.add_paragraph(f"{s['step']}. {s['title']}: {s['detail']}", style="List Number")
        doc.add_heading("Alertas", 2)
        for i in analysis["issues"]: doc.add_paragraph(f"[{i['category']}] {i['title']}: {i['detail']}", style="List Bullet")
        out = io.BytesIO(); doc.save(out)
        return Response(out.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition":"attachment; filename=vaelith-analise.docx"})
    if fmt == "pdf":
        out = io.BytesIO(); styles = getSampleStyleSheet(); story = [Paragraph(title, styles["Title"]), Paragraph("Soluções em Engenharia", styles["Heading2"]), Spacer(1, 6*mm), Paragraph("Conclusão", styles["Heading2"]), Paragraph(analysis["conclusion"], styles["BodyText"]), Spacer(1, 4*mm), Paragraph("Linha de raciocínio", styles["Heading2"])]
        for s in analysis["trace"]: story += [Paragraph(f"<b>{s['step']}. {s['title']}:</b> {s['detail']}", styles["BodyText"]), Spacer(1, 2*mm)]
        story += [Spacer(1, 3*mm), Paragraph("Alertas", styles["Heading2"])]
        table_data = [["Categoria","Título","Detalhe"]] + [[i["category"],i["title"],i["detail"]] for i in analysis["issues"]]
        table = Table(table_data, repeatRows=1, colWidths=[35*mm,50*mm,95*mm]); table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#CFFF47")),("GRID",(0,0),(-1,-1),.3,colors.grey),("VALIGN",(0,0),(-1,-1),"TOP"),("FONTSIZE",(0,0),(-1,-1),7)])); story.append(table)
        SimpleDocTemplate(out, pagesize=A4, leftMargin=15*mm, rightMargin=15*mm, topMargin=15*mm, bottomMargin=15*mm).build(story)
        return Response(out.getvalue(), media_type="application/pdf", headers={"Content-Disposition":"attachment; filename=vaelith-analise.pdf"})
    raise HTTPException(400, "Formato não suportado")


@app.get("/")
def landing(): return FileResponse(BASE / "index.html")
@app.get("/login")
def login_page(): return FileResponse(BASE / "login.html")
@app.get("/app")
def app_page(): return FileResponse(BASE / "app.html")

app.mount("/assets", StaticFiles(directory=BASE / "assets"), name="assets")

if __name__ == "__main__":
    uvicorn.run("server:app", host="0.0.0.0", port=int(os.getenv("PORT", "8080")), reload=False)
