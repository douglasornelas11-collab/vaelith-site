from pathlib import Path
from openpyxl import Workbook
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

ROOT=Path(__file__).resolve().parent/'sample_files'
ROOT.mkdir(exist_ok=True)

def ifc_model(path:Path, elements):
    lines=[]
    def add(s):
        lines.append(f"#{len(lines)+1}={s};")
        return len(lines)
    person=add("IFCPERSON($,$,'VAELITH',$,$,$,$,$)")
    org=add("IFCORGANIZATION($,'VAELITH LABS',$,$,$)")
    pao=add(f"IFCPERSONANDORGANIZATION(#{person},#{org},$)")
    app=add(f"IFCAPPLICATION(#{org},'4.0','VAELITH','VAELITH')")
    oh=add(f"IFCOWNERHISTORY(#{pao},#{app},$,.ADDED.,$,$,$,1721952000)")
    origin=add("IFCCARTESIANPOINT((0.,0.,0.))")
    axis=add(f"IFCAXIS2PLACEMENT3D(#{origin},$,$)")
    context=add(f"IFCGEOMETRICREPRESENTATIONCONTEXT($,'Model',3,1.E-05,#{axis},$)")
    unit=add("IFCSIUNIT(*,.LENGTHUNIT.,$,.METRE.)")
    units=add(f"IFCUNITASSIGNMENT((#{unit}))")
    project=add(f"IFCPROJECT('0J$Q4qA4L7AvY8F3u1P000',#{oh},'Projeto VAELITH',$,$,$,$,(#{context}),#{units})")
    site_pt=add("IFCCARTESIANPOINT((0.,0.,0.))")
    site_ax=add(f"IFCAXIS2PLACEMENT3D(#{site_pt},$,$)")
    site_pl=add(f"IFCLOCALPLACEMENT($,#{site_ax})")
    site=add(f"IFCSITE('0J$Q4qA4L7AvY8F3u1P001',#{oh},'Site',$,$,#{site_pl},$,$,.ELEMENT.,$,$,$,$,$)")
    bpt=add("IFCCARTESIANPOINT((0.,0.,0.))")
    bax=add(f"IFCAXIS2PLACEMENT3D(#{bpt},$,$)")
    bpl=add(f"IFCLOCALPLACEMENT(#{site_pl},#{bax})")
    building=add(f"IFCBUILDING('0J$Q4qA4L7AvY8F3u1P002',#{oh},'Edificio',$,$,#{bpl},$,$,.ELEMENT.,$,$,$)")
    spt=add("IFCCARTESIANPOINT((0.,0.,0.))")
    sax=add(f"IFCAXIS2PLACEMENT3D(#{spt},$,$)")
    spl=add(f"IFCLOCALPLACEMENT(#{bpl},#{sax})")
    storey=add(f"IFCBUILDINGSTOREY('0J$Q4qA4L7AvY8F3u1P003',#{oh},'Pavimento 01',$,$,#{spl},$,$,.ELEMENT.,0.)")
    add(f"IFCRELAGGREGATES('0J$Q4qA4L7AvY8F3u1P004',#{oh},$,$,#{project},(#{site}))")
    add(f"IFCRELAGGREGATES('0J$Q4qA4L7AvY8F3u1P005',#{oh},$,$,#{site},(#{building}))")
    add(f"IFCRELAGGREGATES('0J$Q4qA4L7AvY8F3u1P006',#{oh},$,$,#{building},(#{storey}))")
    product_ids=[]
    for idx,e in enumerate(elements):
        x,y,z=e['pos']; lx,ly,lz=e['size']; guid=e['guid']; name=e['name']; typ=e['type']
        p=add(f"IFCCARTESIANPOINT(({x:.3f},{y:.3f},{z:.3f}))")
        ax=add(f"IFCAXIS2PLACEMENT3D(#{p},$,$)")
        pl=add(f"IFCLOCALPLACEMENT(#{spl},#{ax})")
        p2=add("IFCCARTESIANPOINT((0.,0.))")
        a2=add(f"IFCAXIS2PLACEMENT2D(#{p2},$)")
        prof=add(f"IFCRECTANGLEPROFILEDEF(.AREA.,$ ,#{a2},{lx:.3f},{ly:.3f})")
        dirz=add("IFCDIRECTION((0.,0.,1.))")
        ep=add("IFCCARTESIANPOINT((0.,0.,0.))")
        ea=add(f"IFCAXIS2PLACEMENT3D(#{ep},$,$)")
        solid=add(f"IFCEXTRUDEDAREASOLID(#{prof},#{ea},#{dirz},{lz:.3f})")
        rep=add(f"IFCSHAPEREPRESENTATION(#{context},'Body','SweptSolid',(#{solid}))")
        shape=add(f"IFCPRODUCTDEFINITIONSHAPE($,$,(#{rep}))")
        if typ=='IFCWALL': prod=add(f"IFCWALL('{guid}',#{oh},'{name}',$,$,#{pl},#{shape},$,.NOTDEFINED.)")
        elif typ=='IFCDOOR': prod=add(f"IFCDOOR('{guid}',#{oh},'{name}',$,$,#{pl},#{shape},$,{lz:.3f},{lx:.3f},.DOOR.,.SINGLE_SWING_LEFT.,$)")
        elif typ=='IFCBEAM': prod=add(f"IFCBEAM('{guid}',#{oh},'{name}',$,$,#{pl},#{shape},$,.NOTDEFINED.)")
        elif typ=='IFCCOLUMN': prod=add(f"IFCCOLUMN('{guid}',#{oh},'{name}',$,$,#{pl},#{shape},$,.NOTDEFINED.)")
        else: prod=add(f"IFCBUILDINGELEMENTPROXY('{guid}',#{oh},'{name}',$,$,#{pl},#{shape},$,.NOTDEFINED.)")
        product_ids.append(prod)
    add(f"IFCRELCONTAINEDINSPATIALSTRUCTURE('0J$Q4qA4L7AvY8F3u1P099',#{oh},$,$,({','.join('#'+str(i) for i in product_ids)}),#{storey})")
    content="ISO-10303-21;\nHEADER;\nFILE_DESCRIPTION(('ViewDefinition [CoordinationView]'),'2;1');\nFILE_NAME('sample.ifc','2026-07-26T00:00:00',('VAELITH'),('VAELITH'),'VAELITH','VAELITH','');\nFILE_SCHEMA(('IFC4'));\nENDSEC;\nDATA;\n"+'\n'.join(lines)+"\nENDSEC;\nEND-ISO-10303-21;\n"
    path.write_text(content,encoding='utf-8')

ifc_model(ROOT/'ARQ_R00.ifc',[
    {'type':'IFCWALL','guid':'0J$Q4qA4L7AvY8F3u1W001','name':'Parede sala','pos':(0,0,0),'size':(5,.2,3)},
    {'type':'IFCDOOR','guid':'0J$Q4qA4L7AvY8F3u1D001','name':'Porta P-034','pos':(1.0,-.02,0),'size':(.9,.25,2.1)},
])
ifc_model(ROOT/'ARQ_R01.ifc',[
    {'type':'IFCWALL','guid':'0J$Q4qA4L7AvY8F3u1W001','name':'Parede sala revisada','pos':(0,0,0),'size':(5.5,.2,3)},
    {'type':'IFCDOOR','guid':'0J$Q4qA4L7AvY8F3u1D001','name':'Porta P-034 deslocada','pos':(3.0,-.02,0),'size':(.9,.25,2.1)},
    {'type':'IFCWALL','guid':'0J$Q4qA4L7AvY8F3u1W002','name':'Fechamento do vao','pos':(1.0,0,0),'size':(.9,.2,2.1)},
])
ifc_model(ROOT/'EST_R01.ifc',[
    {'type':'IFCBEAM','guid':'0J$Q4qA4L7AvY8F3u1B001','name':'Viga V-12','pos':(2.6,-.3,2.0),'size':(1.5,.6,.5)},
    {'type':'IFCCOLUMN','guid':'0J$Q4qA4L7AvY8F3u1C001','name':'Pilar P-04','pos':(4.7,-.25,0),'size':(.5,.5,3.2)},
])

wb=Workbook();ws=wb.active;ws.title='Orçamento';ws.append(['Descrição','Quantidade','Unidade','Preço unitário','Valor total']);rows=[
('Abertura de vão em alvenaria',1,'un',420,420),('Fechamento de vão em alvenaria',1,'un',360,360),('Remanejamento de interruptor',1,'ponto',190,190),('Recomposição de reboco e pintura',4.2,'m²',68,285.6),('Reinstalação de porta',1,'un',180,180)]
for r in rows:ws.append(r)
wb.save(ROOT/'ORC_R01.xlsx')
wb=Workbook();ws=wb.active;ws.title='Cronograma';ws.append(['Atividade','Duração','Início','Fim','Predecessora']);rows=[('Alvenaria e vãos',2,'01/08/2026','02/08/2026','Aprovação'),('Instalações elétricas',1,'03/08/2026','03/08/2026','Alvenaria e vãos'),('Reboco e pintura',2,'04/08/2026','05/08/2026','Instalações elétricas')]
for r in rows:ws.append(r)
wb.save(ROOT/'CRONO_R01.xlsx')

doc=Document();doc.add_heading('Memorial de alteração R01',0);doc.add_paragraph('A porta P-034 será deslocada para a parede lateral da Sala de Reuniões 02. O interruptor existente deve ser remanejado e o vão original fechado.');doc.save(ROOT/'ARQ_MEMORIAL_R01.docx')

c=canvas.Canvas(str(ROOT/'ARQ_NOTA_R01.pdf'),pagesize=A4);c.drawString(70,780,'Revisão R01 - Porta P-034');c.drawString(70,750,'Deslocar a porta para a parede lateral e revisar o ponto elétrico.');c.save()
print('samples:',*[p.name for p in ROOT.iterdir()])
